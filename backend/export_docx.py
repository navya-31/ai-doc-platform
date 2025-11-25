# export_docx.py
from docx import Document
from io import BytesIO

def build_docx_bytes(sections):
    """
    sections: list of dicts with 'title' and 'content'
    returns bytes buffer
    """
    doc = Document()
    for sec in sections:
        if sec.get("title"):
            doc.add_heading(sec.get("title"), level=1)
        if sec.get("content"):
            # keep paragraphs intact roughly
            for para in sec.get("content").split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
